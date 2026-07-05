from typing import Optional

from docx import Document


def extract_text_from_docx(docx_file) -> Optional[str]:
    """Extract text from a .docx file-like object, including paragraphs and table cells."""
    document = Document(docx_file)
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    return text or None
